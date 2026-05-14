"""Gateway routes for Dreamy tab — workflow.json CRUD."""

import asyncio
import csv
import json
import logging
import mimetypes
import shutil
import sys
import tempfile
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import Response
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import BaseModel

from src.config.app_config import get_app_config
from src.config.paths import get_paths
from src.models.factory import create_chat_model
from src.models.router import ModelRouter

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["dreamy"])

_REPO_OVERVIEW_PROMPT = """Conduct an indepth analysis of the mounted folder at /mnt/user-data/mounted and write a complete report on all critical files and main features of the mounted folder.
for specific files, there is a mirror repo created within "mnt/user-data/outputs/.docs/" as a identical mirrored(markdownfiles) of the mounted folder."""
_REPO_OVERVIEW_MODEL_TIMEOUT_SECONDS = 45.0


@dataclass
class _RepoOverviewRefreshJob:
    job_id: str
    thread_id: str
    status: str
    created_at: str
    started_at: str | None = None
    finished_at: str | None = None
    error: str | None = None
    output_virtual_path: str | None = None
    trigger: str = "manual"


_REPO_OVERVIEW_JOBS: dict[str, _RepoOverviewRefreshJob] = {}
_REPO_OVERVIEW_JOB_BY_THREAD: dict[str, str] = {}
_REPO_OVERVIEW_TASKS: dict[str, asyncio.Task[None]] = {}

_SKIP_DIR_NAMES = {
    ".git",
    ".hg",
    ".svn",
    "__pycache__",
    ".mypy_cache",
    ".pytest_cache",
    ".ruff_cache",
    ".next",
    ".turbo",
    "node_modules",
    "dist",
    "build",
    "coverage",
    ".idea",
    ".vscode",
}

_TEXT_EXTENSIONS = {
    ".md",
    ".txt",
    ".rst",
    ".py",
    ".ts",
    ".tsx",
    ".js",
    ".jsx",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".sh",
    ".bash",
    ".zsh",
    ".sql",
    ".csv",
    ".tsv",
    ".xml",
    ".html",
    ".css",
    ".scss",
    ".go",
    ".rs",
    ".java",
    ".kt",
    ".swift",
    ".c",
    ".h",
    ".cpp",
    ".hpp",
}


def _escape_md_cell(value: object) -> str:
    text = str(value) if value is not None else ""
    return text.replace("|", "\\|").replace("\n", "<br>")


def _rows_to_markdown_table(rows: list[list[object]], *, max_rows: int = 200) -> str:
    if not rows:
        return "_No rows found._"
    header = rows[0]
    body = rows[1:]
    shown = body[:max_rows]
    lines = [
        "| " + " | ".join(_escape_md_cell(c) for c in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in shown:
        padded = list(row) + [""] * max(0, len(header) - len(row))
        lines.append("| " + " | ".join(_escape_md_cell(c) for c in padded[: len(header)]) + " |")
    if len(body) > max_rows:
        lines.append("")
        lines.append(f"_Truncated to first {max_rows} data rows out of {len(body)}._")
    return "\n".join(lines)


def _read_pdf_text(path: Path) -> tuple[str | None, str | None]:
    try:
        from pypdf import PdfReader
    except Exception:
        return None, "pypdf-not-installed"
    try:
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            pages.append(f"## Page {i + 1}\n\n{text or '_No extractable text on this page._'}")
        return "\n\n".join(pages), None
    except Exception as exc:
        return None, f"pdf-read-failed: {exc}"


def _read_docx_text(path: Path) -> tuple[str | None, str | None]:
    try:
        from docx import Document
    except Exception:
        return None, "python-docx-not-installed"
    try:
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)
        for idx, table in enumerate(doc.tables):
            rows: list[list[object]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                parts.append(f"## Table {idx + 1}\n\n{_rows_to_markdown_table(rows)}")
        return "\n\n".join(parts), None
    except Exception as exc:
        return None, f"docx-read-failed: {exc}"


def _read_image_text(path: Path) -> tuple[str | None, str | None]:
    try:
        from PIL import Image
    except Exception:
        return None, "pillow-not-installed"
    try:
        import pytesseract
    except Exception:
        return None, "pytesseract-not-installed"
    try:
        with Image.open(path) as img:
            text = pytesseract.image_to_string(img).strip()
        if not text:
            return "_No OCR text extracted from image._", None
        return text, None
    except Exception as exc:
        return None, f"image-ocr-failed: {exc}"


def _read_csv_like(path: Path, delimiter: str = ",") -> tuple[str | None, str | None]:
    try:
        with path.open("r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f, delimiter=delimiter)
            rows = [row for row in reader]
        return _rows_to_markdown_table(rows), None
    except Exception as exc:
        return None, f"tabular-read-failed: {exc}"


def _read_xlsx(path: Path) -> tuple[str | None, str | None]:
    try:
        from openpyxl import load_workbook
    except Exception:
        return None, "openpyxl-not-installed"
    try:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
        sections: list[str] = []
        for sheet in wb.worksheets:
            rows: list[list[object]] = []
            for row in sheet.iter_rows(values_only=True):
                rows.append(list(row))
            sections.append(f"## Sheet: {sheet.title}\n\n{_rows_to_markdown_table(rows)}")
        return "\n\n".join(sections), None
    except Exception as exc:
        return None, f"xlsx-read-failed: {exc}"


def _extract_to_markdown(path: Path) -> tuple[str | None, str | None, str | None]:
    """Return (markdown_content, source_kind, error_reason)."""
    suffix = path.suffix.lower()
    if suffix in {".csv"}:
        content, err = _read_csv_like(path, ",")
        return content, "csv", err
    if suffix in {".tsv"}:
        content, err = _read_csv_like(path, "\t")
        return content, "tsv", err
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        content, err = _read_xlsx(path)
        return content, "xlsx", err
    if suffix == ".pdf":
        content, err = _read_pdf_text(path)
        return content, "pdf", err
    if suffix == ".docx":
        content, err = _read_docx_text(path)
        return content, "docx", err
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}:
        content, err = _read_image_text(path)
        return content, "image_ocr", err
    if _looks_like_text(path):
        try:
            return path.read_text(encoding="utf-8"), "text", None
        except UnicodeDecodeError:
            return None, None, "utf8-decode-failed"
        except Exception as exc:
            return None, None, f"read-failed: {exc}"
    return None, None, "non-text-or-binary"


def _looks_like_text(path: Path) -> bool:
    suffix = path.suffix.lower()
    if suffix in _TEXT_EXTENSIONS:
        return True
    guessed, _ = mimetypes.guess_type(path.name)
    if guessed and guessed.startswith("text/"):
        return True
    try:
        chunk = path.read_bytes()[:4096]
    except Exception:
        return False
    if b"\x00" in chunk:
        return False
    try:
        chunk.decode("utf-8")
        return True
    except UnicodeDecodeError:
        return False


def _guess_fence_language(path: Path) -> str:
    ext = path.suffix.lower().lstrip(".")
    if not ext:
        return "text"
    if ext == "yml":
        return "yaml"
    if ext == "md":
        return "markdown"
    if ext in {"tsx", "jsx"}:
        return ext
    return ext


def _to_virtual_outputs(path: Path, thread_id: str) -> str:
    outputs_dir = get_paths().sandbox_outputs_dir(thread_id).resolve()
    rel = path.resolve().relative_to(outputs_dir).as_posix()
    return f"/mnt/user-data/outputs/{rel}"


def _summarize_markdown_corpus(docs_root: Path) -> dict[str, object]:
    """Build a deep, Python-only repo summary from generated markdown mirrors."""
    docs: list[Path] = []
    total_chars = 0
    keyword_hits: dict[str, int] = {
        "api": 0,
        "database": 0,
        "auth": 0,
        "config": 0,
        "pipeline": 0,
        "workflow": 0,
        "test": 0,
    }
    dir_signal_counts: dict[str, int] = {}
    high_signal_files: list[str] = []

    for p in sorted(docs_root.rglob("*.md")):
        if p.name in {
            "index.md",
            "failed_files.md",
            "created_files.md",
            "file_catalog.md",
            "directory_tree.md",
            "repo_overview.md",
        }:
            continue
        docs.append(p)
        rel = p.relative_to(docs_root).as_posix()
        top = rel.split("/", 1)[0] if "/" in rel else "(root)"
        dir_signal_counts[top] = dir_signal_counts.get(top, 0) + 1
        if any(
            token in rel.lower()
            for token in (
                "readme",
                "main.",
                "app.",
                "server.",
                "dockerfile",
                "pyproject",
                "package.json",
                "config.",
            )
        ):
            high_signal_files.append(rel)
        try:
            text = p.read_text(encoding="utf-8")
        except Exception:
            continue
        total_chars += len(text)
        lowered = text.lower()
        for keyword in keyword_hits:
            keyword_hits[keyword] += lowered.count(keyword)

    avg_chars = int(total_chars / len(docs)) if docs else 0
    strongest_keywords = sorted(keyword_hits.items(), key=lambda item: (-item[1], item[0]))
    top_dirs = sorted(dir_signal_counts.items(), key=lambda item: (-item[1], item[0]))
    return {
        "docs_count": len(docs),
        "total_chars": total_chars,
        "avg_chars": avg_chars,
        "keyword_hits": keyword_hits,
        "strongest_keywords": strongest_keywords,
        "top_dirs": top_dirs,
        "high_signal_files": sorted(set(high_signal_files))[:200],
    }


def _workflow_path(thread_id: str) -> Path:
    return get_paths().sandbox_outputs_dir(thread_id) / "workflow.json"


def _mount_config_path(thread_id: str) -> Path:
    return get_paths().sandbox_user_data_dir(thread_id) / "dreamy_mount.json"


def _job_to_response(job: _RepoOverviewRefreshJob) -> "RepoOverviewRefreshStatusResponse":
    return RepoOverviewRefreshStatusResponse(
        job_id=job.job_id,
        thread_id=job.thread_id,
        status=job.status,
        created_at=job.created_at,
        started_at=job.started_at,
        finished_at=job.finished_at,
        error=job.error,
        output_virtual_path=job.output_virtual_path,
    )


def _read_text_if_exists(path: Path, *, max_chars: int = 80000) -> str:
    if not path.exists() or not path.is_file():
        return ""
    try:
        raw = path.read_text(encoding="utf-8")
        if len(raw) > max_chars:
            return raw[:max_chars] + "\n\n[...truncated...]"
        return raw
    except Exception:
        return ""


def _build_repo_overview_refresh_prompt(source_root: Path, docs_root: Path) -> str:
    index_text = _read_text_if_exists(docs_root / "index.md")
    tree_text = _read_text_if_exists(docs_root / "directory_tree.md")
    catalog_text = _read_text_if_exists(docs_root / "file_catalog.md", max_chars=120000)
    failed_text = _read_text_if_exists(docs_root / "failed_files.md", max_chars=40000)
    return "\n\n".join(
        [
            _REPO_OVERVIEW_PROMPT,
            f"Mounted root (host path): {source_root}",
            "Use the mirrored markdown docs under /mnt/user-data/outputs/.docs as your primary evidence base.",
            "Output only valid markdown for repo_overview.md.",
            "Cover architecture, critical files, key features, execution flow, risks, and recommended next reading order.",
            "Be concrete and cite file paths from the mirror.",
            "Context from generated docs:",
            "## index.md",
            index_text or "(missing)",
            "## directory_tree.md",
            tree_text or "(missing)",
            "## file_catalog.md",
            catalog_text or "(missing)",
            "## failed_files.md",
            failed_text or "(missing)",
        ]
    )


async def _run_repo_overview_refresh_job(job_id: str, source_root: Path, docs_root: Path) -> None:
    job = _REPO_OVERVIEW_JOBS[job_id]
    job.status = "running"
    job.started_at = datetime.now(UTC).isoformat()
    try:
        if not docs_root.exists() or not docs_root.is_dir():
            raise RuntimeError("Docs mirror not found. Run /analyse first.")
        if not (docs_root / "index.md").exists():
            raise RuntimeError("Docs index not found. Run /analyse first.")

        prompt = _build_repo_overview_refresh_prompt(source_root, docs_root)
        router = ModelRouter()
        primary_model = router.resolve("planner")
        configured_models = [m.name for m in get_app_config().models]
        ordered_models: list[str] = []
        if primary_model:
            ordered_models.append(primary_model)
        for name in configured_models:
            if name not in ordered_models:
                ordered_models.append(name)

        response = None
        last_exc: Exception | None = None
        for candidate in ordered_models:
            try:
                model = create_chat_model(
                    name=candidate,
                    thinking_enabled=True,
                    reasoning_effort="high",
                )
                response = await asyncio.wait_for(
                    asyncio.to_thread(
                        model.invoke,
                        [
                            SystemMessage(
                                content=(
                                    "You are an expert software architect and reviewer. "
                                    "Generate a detailed, practical repository analysis report in markdown only."
                                )
                            ),
                            HumanMessage(content=prompt),
                        ],
                    ),
                    timeout=_REPO_OVERVIEW_MODEL_TIMEOUT_SECONDS,
                )
                break
            except Exception as exc:
                last_exc = exc
                logger.warning(
                    "Repo overview refresh model '%s' failed for thread %s: %s",
                    candidate,
                    job.thread_id,
                    exc,
                )
        if response is None:
            raise RuntimeError(f"All configured models failed during repo overview refresh. Last error: {last_exc}")
        content = getattr(response, "content", "")
        if isinstance(content, list):
            text_parts: list[str] = []
            for part in content:
                if isinstance(part, dict) and part.get("type") == "text":
                    text_parts.append(str(part.get("text", "")))
                else:
                    text_parts.append(str(part))
            content_text = "\n".join(text_parts).strip()
        else:
            content_text = str(content).strip()
        if not content_text:
            raise RuntimeError("Model returned empty repo overview content.")

        target = docs_root / "repo_overview.md"
        backup = docs_root / "repo_overview.previous.md"
        if target.exists():
            shutil.copy2(target, backup)
        tmp = docs_root / f".repo_overview.{job_id}.tmp.md"
        tmp.write_text(content_text + "\n", encoding="utf-8")
        tmp.replace(target)
        job.output_virtual_path = f"/mnt/user-data/outputs/.docs/{target.name}"
        job.status = "succeeded"
    except Exception as exc:
        logger.exception("Repo overview refresh failed for thread %s: %s", job.thread_id, exc)
        job.status = "failed"
        job.error = str(exc)
    finally:
        job.finished_at = datetime.now(UTC).isoformat()
        _REPO_OVERVIEW_TASKS.pop(job_id, None)
        current = _REPO_OVERVIEW_JOB_BY_THREAD.get(job.thread_id)
        if current == job_id:
            _REPO_OVERVIEW_JOB_BY_THREAD.pop(job.thread_id, None)


async def _enqueue_repo_overview_refresh(
    thread_id: str,
    source_root: Path,
    docs_root: Path,
    *,
    trigger: str,
) -> "RepoOverviewRefreshStartResponse":
    running_job_id = _REPO_OVERVIEW_JOB_BY_THREAD.get(thread_id)
    if running_job_id:
        running = _REPO_OVERVIEW_JOBS.get(running_job_id)
        if running and running.status in {"queued", "running"}:
            return RepoOverviewRefreshStartResponse(
                job_id=running.job_id,
                status=running.status,
                already_running=True,
            )

    job_id = str(uuid4())
    job = _RepoOverviewRefreshJob(
        job_id=job_id,
        thread_id=thread_id,
        status="queued",
        created_at=datetime.now(UTC).isoformat(),
        trigger=trigger,
    )
    _REPO_OVERVIEW_JOBS[job_id] = job
    _REPO_OVERVIEW_JOB_BY_THREAD[thread_id] = job_id
    task = asyncio.create_task(_run_repo_overview_refresh_job(job_id, source_root, docs_root))
    _REPO_OVERVIEW_TASKS[job_id] = task
    return RepoOverviewRefreshStartResponse(job_id=job_id, status=job.status, already_running=False)


def _maybe_migrate_v1(data: dict) -> dict:
    """In-memory migration from v1 DAG schema to v2 steps schema. Does not rewrite the file."""
    if data.get("version") != "1":
        return data
    ts = data.get("task_source", {})
    es = data.get("execution_state", {})
    phase = es.get("phase", "design")
    if phase == "approval":
        phase = "awaiting_approval"
    return {
        "version": "2",
        "thread_id": data.get("thread_id"),
        "created_at": data.get("created_at"),
        "data_source": {
            "type": ts.get("type", "inline"),
            "filename": ts.get("filename", ""),
            "total_rows": ts.get("total_tasks", 0),
            "fields": ts.get("fields", []),
            "sample_rows": ts.get("sample_tasks", []),
        },
        "steps": [],
        "execution_state": {
            "phase": phase,
            "current_row_index": es.get("current_task_index", 0),
            "current_step_id": es.get("active_node_id"),
            "total_rows": es.get("total_tasks", 0),
            "poc_results": [],
            "seconds_per_row_estimate": None,
            "estimated_completion_iso": es.get("estimated_completion_iso"),
            "started_at": es.get("started_at"),
        },
    }


@router.get("/threads/{thread_id}/dreamy/workflow")
async def get_workflow(thread_id: str) -> dict:
    """Return the current workflow.json for a Dreamy thread (auto-migrates v1 → v2 in memory)."""
    path = _workflow_path(thread_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="workflow.json not found")
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return _maybe_migrate_v1(data)
    except Exception as exc:
        logger.error("Failed to read workflow.json for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to read workflow.json") from exc


class WorkflowPatchRequest(BaseModel):
    workflow: dict


@router.patch("/threads/{thread_id}/dreamy/workflow")
async def patch_workflow(thread_id: str, req: WorkflowPatchRequest) -> dict:
    """Persist user edits to workflow.json from the node editor."""
    path = _workflow_path(thread_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        path.write_text(json.dumps(req.workflow, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as exc:
        logger.error("Failed to write workflow.json for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to write workflow.json") from exc
    return {"success": True}


class MountFolderRequest(BaseModel):
    path: str


class AnalyseResponse(BaseModel):
    source_root: str
    output_root: str
    generated_docs: int
    skipped_non_text: int
    failed: int
    index_virtual_path: str
    failed_manifest_virtual_path: str
    failed_files: list[dict[str, str]]
    repo_overview_refresh_job_id: str | None = None


class AnalyseStatusResponse(BaseModel):
    staged_available: bool
    docs_root_virtual_path: str
    index_virtual_path: str | None = None
    failed_manifest_virtual_path: str | None = None


class PublishDocsResponse(BaseModel):
    source_root: str
    destination_root: str
    copied_files: int
    overwritten_files: int
    index_virtual_path: str


class RepoOverviewRefreshStartResponse(BaseModel):
    job_id: str
    status: str
    already_running: bool = False


class RepoOverviewRefreshStatusResponse(BaseModel):
    job_id: str
    thread_id: str
    status: str
    created_at: str
    started_at: str | None = None
    finished_at: str | None = None
    error: str | None = None
    output_virtual_path: str | None = None


@router.get("/threads/{thread_id}/analyse/status", response_model=AnalyseStatusResponse)
async def get_analyse_status(thread_id: str) -> AnalyseStatusResponse:
    outputs_root = get_paths().sandbox_outputs_dir(thread_id).resolve()
    docs_root = outputs_root / ".docs"
    index_path = docs_root / "index.md"
    failed_manifest = docs_root / "failed_files.md"
    return AnalyseStatusResponse(
        staged_available=docs_root.exists() and docs_root.is_dir() and index_path.exists(),
        docs_root_virtual_path="/mnt/user-data/outputs/.docs",
        index_virtual_path=_to_virtual_outputs(index_path, thread_id) if index_path.exists() else None,
        failed_manifest_virtual_path=_to_virtual_outputs(failed_manifest, thread_id) if failed_manifest.exists() else None,
    )


@router.post("/threads/{thread_id}/analyse/repo-overview-refresh", response_model=RepoOverviewRefreshStartResponse)
async def start_repo_overview_refresh(thread_id: str) -> RepoOverviewRefreshStartResponse:
    config_path = _mount_config_path(thread_id)
    if not config_path.exists():
        raise HTTPException(status_code=400, detail="No mounted folder configured for this thread.")
    try:
        data = json.loads(config_path.read_text(encoding="utf-8"))
        mounted_path_str = (data.get("path") or "").strip()
    except Exception as exc:
        logger.error("Failed to read dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to read mounted folder config") from exc
    if not mounted_path_str:
        raise HTTPException(status_code=400, detail="Mounted folder path is empty.")

    source_root = Path(mounted_path_str).expanduser().resolve()
    if not source_root.exists() or not source_root.is_dir():
        raise HTTPException(status_code=400, detail="Mounted folder does not exist or is not a directory.")

    docs_root = get_paths().sandbox_outputs_dir(thread_id).resolve() / ".docs"
    if not docs_root.exists() or not docs_root.is_dir() or not (docs_root / "index.md").exists():
        raise HTTPException(status_code=400, detail="Staged docs index not found. Run /analyse first.")

    return await _enqueue_repo_overview_refresh(
        thread_id=thread_id,
        source_root=source_root,
        docs_root=docs_root,
        trigger="manual",
    )


@router.get(
    "/threads/{thread_id}/analyse/repo-overview-refresh/{job_id}",
    response_model=RepoOverviewRefreshStatusResponse,
)
async def get_repo_overview_refresh_status(thread_id: str, job_id: str) -> RepoOverviewRefreshStatusResponse:
    job = _REPO_OVERVIEW_JOBS.get(job_id)
    if job is None or job.thread_id != thread_id:
        raise HTTPException(status_code=404, detail="Repo overview refresh job not found.")
    return _job_to_response(job)


@router.post("/threads/{thread_id}/analyse", response_model=AnalyseResponse)
async def run_analyse(thread_id: str) -> AnalyseResponse:
    """Deterministically mirror mounted repository files into /outputs/.docs as Markdown."""
    config_path = _mount_config_path(thread_id)
    if not config_path.exists():
        raise HTTPException(status_code=400, detail="No mounted folder configured for this thread.")
    try:
        data = json.loads(config_path.read_text(encoding="utf-8"))
        mounted_path_str = (data.get("path") or "").strip()
    except Exception as exc:
        logger.error("Failed to read dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to read mounted folder config") from exc
    if not mounted_path_str:
        raise HTTPException(status_code=400, detail="Mounted folder path is empty.")

    source_root = Path(mounted_path_str).expanduser().resolve()
    if not source_root.exists() or not source_root.is_dir():
        raise HTTPException(status_code=400, detail="Mounted folder does not exist or is not a directory.")

    outputs_root = get_paths().sandbox_outputs_dir(thread_id).resolve()
    docs_root = outputs_root / ".docs"
    docs_root.mkdir(parents=True, exist_ok=True)

    generated_docs = 0
    skipped_non_text = 0
    failed_files: list[dict[str, str]] = []
    converted_files: list[dict[str, str]] = []
    extension_counts: dict[str, int] = {}
    top_level_counts: dict[str, int] = {}
    source_kind_counts: dict[str, int] = {}

    for source_path in sorted(source_root.rglob("*")):
        if not source_path.is_file():
            continue
        rel = source_path.relative_to(source_root)
        if any(part in _SKIP_DIR_NAMES for part in rel.parts):
            continue
        if rel.parts and rel.parts[0] == ".docs":
            continue
        content, source_kind, error_reason = _extract_to_markdown(source_path)
        if content is None:
            if error_reason == "non-text-or-binary":
                skipped_non_text += 1
            failed_files.append({"path": rel.as_posix(), "reason": error_reason or "conversion-failed"})
            continue

        target_path = docs_root / rel.parent / f"{rel.name}.md"
        target_path.parent.mkdir(parents=True, exist_ok=True)
        ext = source_path.suffix.lower() or "(no_ext)"
        extension_counts[ext] = extension_counts.get(ext, 0) + 1
        top = rel.parts[0] if rel.parts else "(root)"
        top_level_counts[top] = top_level_counts.get(top, 0) + 1
        if source_kind == "text":
            language = _guess_fence_language(source_path)
            body = "\n".join(
                [
                    f"# {rel.as_posix()}",
                    "",
                    "## Purpose",
                    f"- Source path: `{rel.as_posix()}`",
                    "- This document mirrors the original file for analysis and retrieval.",
                    "",
                    "## Content",
                    f"```{language}",
                    content,
                    "```",
                    "",
                ]
            )
        else:
            body = "\n".join(
                [
                    f"# {rel.as_posix()}",
                    "",
                    "## Purpose",
                    f"- Source path: `{rel.as_posix()}`",
                    f"- Converted source type: `{source_kind}`",
                    "- This document mirrors extracted content for analysis and retrieval.",
                    "",
                    "## Extracted Content",
                    content,
                    "",
                ]
            )
        try:
            target_path.write_text(body, encoding="utf-8")
            generated_docs += 1
            source_kind_counts[source_kind or "unknown"] = source_kind_counts.get(source_kind or "unknown", 0) + 1
            converted_files.append(
                {
                    "source_path": rel.as_posix(),
                    "source_kind": source_kind or "unknown",
                    "doc_virtual_path": _to_virtual_outputs(target_path, thread_id),
                }
            )
        except Exception as exc:
            failed_files.append({"path": rel.as_posix(), "reason": f"write-failed: {exc}"})

    failed_manifest_path = docs_root / "failed_files.md"
    failed_lines = [
        "# Failed Files During /analyse",
        "",
        f"- Generated at: {datetime.now(UTC).isoformat()}",
        f"- Total failures: {len(failed_files)}",
        "",
    ]
    if failed_files:
        failed_lines.append("| Path | Reason |")
        failed_lines.append("| --- | --- |")
        for item in failed_files:
            failed_lines.append(f"| `{item['path']}` | `{item['reason']}` |")
    else:
        failed_lines.append("No failures.")
    failed_manifest_path.write_text("\n".join(failed_lines) + "\n", encoding="utf-8")

    directory_tree_path = docs_root / "directory_tree.md"
    critical_paths: list[str] = []
    for item in converted_files:
        src = item["source_path"].lower()
        if (
            src.endswith("readme.md")
            or "/src/" in src
            or src.endswith("main.py")
            or src.endswith("app.py")
            or src.endswith("server.py")
            or src.endswith("dockerfile")
            or src.endswith("docker-compose.yml")
            or src.endswith("pyproject.toml")
            or src.endswith("package.json")
            or src.endswith("config.yaml")
            or src.endswith("config.yml")
        ):
            critical_paths.append(item["source_path"])
    critical_paths = sorted(set(critical_paths))[:200]

    tree_lines = [
        "# Directory Tree",
        "",
        f"- Source root: `{source_root}`",
        f"- Total converted files: **{len(converted_files)}**",
        f"- Total failed files: **{len(failed_files)}**",
        "",
        "## Critical Paths For Future Prompts",
    ]
    if critical_paths:
        for path in critical_paths:
            tree_lines.append(f"- `{path}`")
    else:
        tree_lines.append("- No critical paths detected with current heuristics.")
    tree_lines.extend(
        [
            "",
            "## Prompting Context For Future Queries",
            "- Treat this tree as the map for where behavior is implemented.",
            "- Start with critical paths, then use `file_catalog.md` to jump to mirrored markdown docs.",
            "- Cross-reference unresolved files with `failed_files.md` before concluding a gap.",
            "",
            "## Full Tree",
            "",
            "```text",
        ]
    )

    for p in sorted(source_root.rglob("*")):
        rel = p.relative_to(source_root)
        if any(part in _SKIP_DIR_NAMES for part in rel.parts):
            continue
        if rel.parts and rel.parts[0] == ".docs":
            continue
        depth = len(rel.parts) - 1
        prefix = "  " * max(0, depth)
        suffix = "/" if p.is_dir() else ""
        tree_lines.append(f"{prefix}{rel.name}{suffix}")
    tree_lines.append("```")
    directory_tree_path.write_text("\n".join(tree_lines) + "\n", encoding="utf-8")

    created_manifest_path = docs_root / "created_files.md"
    created_lines = [
        "# Created Files During /analyse",
        "",
        f"- Generated at: {datetime.now(UTC).isoformat()}",
        f"- Total created markdown mirrors: **{len(converted_files)}**",
        "",
        "| Source Path | Source Type | Markdown Mirror |",
        "| --- | --- | --- |",
    ]
    for item in converted_files[:4000]:
        created_lines.append(
            f"| `{item['source_path']}` | `{item['source_kind']}` | `{item['doc_virtual_path']}` |"
        )
    if len(converted_files) > 4000:
        created_lines.extend(
            [
                "",
                f"_Truncated list: showing first 4000 of {len(converted_files)} created markdown files._",
            ]
        )
    created_manifest_path.write_text("\n".join(created_lines) + "\n", encoding="utf-8")

    file_catalog_path = docs_root / "file_catalog.md"
    catalog_lines = [
        "# File Catalog",
        "",
        f"- Total converted files: **{len(converted_files)}**",
        "",
        "| Source Path | Type | Doc Path |",
        "| --- | --- | --- |",
    ]
    for item in converted_files[:2000]:
        catalog_lines.append(
            f"| `{item['source_path']}` | `{item['source_kind']}` | `{item['doc_virtual_path']}` |"
        )
    if len(converted_files) > 2000:
        catalog_lines.extend(
            [
                "",
                f"_Truncated catalog output: showing first 2000 of {len(converted_files)} files._",
            ]
        )
    file_catalog_path.write_text("\n".join(catalog_lines) + "\n", encoding="utf-8")

    repo_overview_path = docs_root / "repo_overview.md"
    corpus = _summarize_markdown_corpus(docs_root)
    high_signal_files = sorted(
        {
            item["source_path"]
            for item in converted_files
            if item["source_path"].lower().endswith(
                (
                    "readme.md",
                    "pyproject.toml",
                    "package.json",
                    "dockerfile",
                    "docker-compose.yml",
                    "config.yaml",
                    "config.yml",
                    "requirements.txt",
                )
            )
            or "/src/" in item["source_path"].lower()
        }
    )[:200]

    overview_lines = [
        "# Repository Overview",
        "",
        "## Analysis Header",
        f"- Number of successful files: **{len(converted_files)}**",
        f"- Number of failed files: **{len(failed_files)}**",
        "- By Extension breakdown:",
    ]
    for key, value in sorted(extension_counts.items(), key=lambda item: (-item[1], item[0])):
        overview_lines.append(f"  - `{key}`: {value} files")
    overview_lines.extend(
        [
            "",
            "## Deep Repository Analysis (Python-based)",
            f"- Analysed markdown mirrors: **{int(corpus['docs_count'])}**",
            f"- Total mirror corpus size: **{int(corpus['total_chars'])}** characters",
            f"- Average file narrative size: **{int(corpus['avg_chars'])}** characters",
            "",
            "### Vital / High-Signal Paths",
        ]
    )
    high_signal_from_corpus = corpus["high_signal_files"] if isinstance(corpus["high_signal_files"], list) else []
    merged_high_signal = sorted(set(high_signal_files + [str(x) for x in high_signal_from_corpus]))[:250]
    if merged_high_signal:
        for path in merged_high_signal:
            overview_lines.append(f"- `{path}`")
    else:
        overview_lines.append("- No high-signal files matched current heuristics.")
    overview_lines.extend(
        [
            "",
            "### Key Feature Signals",
        ]
    )
    strongest = corpus["strongest_keywords"] if isinstance(corpus["strongest_keywords"], list) else []
    for entry in strongest[:7]:
        if isinstance(entry, tuple) and len(entry) == 2:
            overview_lines.append(f"- `{entry[0]}` signal mentions: **{entry[1]}**")
    overview_lines.extend(
        [
            "",
            "### What Failed vs What Was Created",
            f"- Created manifest: `{_to_virtual_outputs(created_manifest_path, thread_id)}`",
            f"- Failed manifest: `{_to_virtual_outputs(failed_manifest_path, thread_id)}`",
            "- Use failed manifest first when investigating blind spots or missing behavior.",
            "",
            "### Recommended Prompting Instructions",
            "```text",
            "Start with /mnt/user-data/outputs/.docs/index.md, repo_overview.md, and directory_tree.md.",
            "Then use file_catalog.md to jump to mirrored files for any area you investigate.",
            "When an answer depends on files listed in failed_files.md, explicitly state the uncertainty.",
            "```",
            "",
            "## Executive Summary",
            f"- Repository docs generated: **{len(converted_files)}** mirrored markdown files.",
            f"- Files failed to process: **{len(failed_files)}** (see `failed_files.md`).",
            f"- File-type diversity: **{len(extension_counts)}** extensions, **{len(source_kind_counts)}** conversion kinds.",
        ]
    )

    overview_lines.extend(["", "## Conversion Coverage By Source Kind"])
    for key, value in sorted(source_kind_counts.items(), key=lambda item: (-item[1], item[0])):
        overview_lines.append(f"- `{key}`: **{value}** files")
    overview_lines.extend(
        [
            "",
            "## High-Signal Files (Quick Start)",
        ]
    )
    if high_signal_files:
        for path in high_signal_files:
            overview_lines.append(f"- `{path}`")
    else:
        overview_lines.append("- No high-signal files matched current heuristics.")
    overview_lines.extend(
        [
            "",
            "## In-Depth Structural Evaluation",
            "- Top-level directory concentration helps identify system boundaries and ownership zones.",
            "- Extension distribution helps infer runtime stack, data contracts, and config surface.",
            "- Critical paths in `directory_tree.md` indicate likely entrypoints and orchestrators.",
            "- Use `file_catalog.md` to map any source path to its mirrored markdown for detailed reading.",
            "",
            "## By Top-Level Directory",
        ]
    )
    for key, value in sorted(top_level_counts.items(), key=lambda item: (-item[1], item[0])):
        overview_lines.append(f"- `{key}`: **{value}** files")
    overview_lines.extend(["", "## By Extension"])
    for key, value in sorted(extension_counts.items(), key=lambda item: (-item[1], item[0])):
        overview_lines.append(f"- `{key}`: **{value}** files")
    overview_lines.extend(
        [
            "",
            "## Recommended LLM Prompt Templates",
            "Use these directly in chat for deeper analysis:",
            "",
            "```text",
            "Using /mnt/user-data/outputs/.docs/index.md, /repo_overview.md, /directory_tree.md, and /file_catalog.md,",
            "give me an architecture-level summary of the repository with:",
            "1) core modules, 2) critical execution flow, 3) highest-risk areas, 4) suggested next code-reading order.",
            "```",
            "",
            "```text",
            "From /mnt/user-data/outputs/.docs/failed_files.md and /created_files.md,",
            "identify any analysis blind spots and propose how to close them.",
            "```",
            "",
            "```text",
            "Use /mnt/user-data/outputs/.docs/file_catalog.md to locate files related to <topic>,",
            "then synthesize findings only from the mirrored markdown docs.",
            "```",
        ]
    )
    repo_overview_path.write_text("\n".join(overview_lines) + "\n", encoding="utf-8")

    index_path = docs_root / "index.md"
    index_lines = [
        "# Repository Docs Mirror",
        "",
        f"- Source root: `{source_root}`",
        "- Output root: `/mnt/user-data/outputs/.docs`",
        f"- Generated docs: **{generated_docs}**",
        f"- Skipped non-text files: **{skipped_non_text}**",
        f"- Failures: **{len(failed_files)}**",
        "",
        "## Notes",
        "- This mirror is deterministic and generated by `/analyse` API.",
        "- Prefer consulting these `.docs` files first for follow-up queries in this thread.",
        f"- Overview: `{_to_virtual_outputs(repo_overview_path, thread_id)}`",
        f"- Directory tree: `{_to_virtual_outputs(directory_tree_path, thread_id)}`",
        f"- Created files manifest: `{_to_virtual_outputs(created_manifest_path, thread_id)}`",
        f"- File catalog: `{_to_virtual_outputs(file_catalog_path, thread_id)}`",
        f"- Failure manifest: `{_to_virtual_outputs(failed_manifest_path, thread_id)}`",
        "",
    ]
    index_path.write_text("\n".join(index_lines), encoding="utf-8")

    refresh_job = await _enqueue_repo_overview_refresh(
        thread_id=thread_id,
        source_root=source_root,
        docs_root=docs_root,
        trigger="analyse",
    )

    return AnalyseResponse(
        source_root=str(source_root),
        output_root="/mnt/user-data/outputs/.docs",
        generated_docs=generated_docs,
        skipped_non_text=skipped_non_text,
        failed=len(failed_files),
        index_virtual_path=_to_virtual_outputs(index_path, thread_id),
        failed_manifest_virtual_path=_to_virtual_outputs(failed_manifest_path, thread_id),
        failed_files=failed_files[:200],
        repo_overview_refresh_job_id=refresh_job.job_id,
    )


@router.post("/threads/{thread_id}/publishdocs", response_model=PublishDocsResponse)
async def publish_docs(thread_id: str) -> PublishDocsResponse:
    config_path = _mount_config_path(thread_id)
    if not config_path.exists():
        raise HTTPException(status_code=400, detail="No mounted folder configured for this thread.")
    try:
        data = json.loads(config_path.read_text(encoding="utf-8"))
        mounted_path_str = (data.get("path") or "").strip()
    except Exception as exc:
        logger.error("Failed to read dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to read mounted folder config") from exc
    if not mounted_path_str:
        raise HTTPException(status_code=400, detail="Mounted folder path is empty.")

    source_root = get_paths().sandbox_outputs_dir(thread_id).resolve() / ".docs"
    if not source_root.exists() or not source_root.is_dir():
        raise HTTPException(status_code=400, detail="Staged docs not found at /mnt/user-data/outputs/.docs.")
    if not (source_root / "index.md").exists():
        raise HTTPException(status_code=400, detail="Staged docs index not found. Run /analyse first.")

    destination_root = Path(mounted_path_str).expanduser().resolve() / ".docs"
    destination_root.mkdir(parents=True, exist_ok=True)

    copied_files = 0
    overwritten_files = 0
    for src in sorted(source_root.rglob("*")):
        if not src.is_file():
            continue
        rel = src.relative_to(source_root)
        dst = destination_root / rel
        dst.parent.mkdir(parents=True, exist_ok=True)
        if dst.exists():
            overwritten_files += 1
        shutil.copy2(src, dst)
        copied_files += 1

    return PublishDocsResponse(
        source_root="/mnt/user-data/outputs/.docs",
        destination_root="/mnt/user-data/mounted/.docs",
        copied_files=copied_files,
        overwritten_files=overwritten_files,
        index_virtual_path="/mnt/user-data/mounted/.docs/index.md",
    )


@router.get("/threads/{thread_id}/dreamy/mount-folder")
async def get_mount_folder(thread_id: str) -> dict:
    path = _mount_config_path(thread_id)
    if not path.exists():
        return {"path": None}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        logger.error("Failed to read dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to read mounted folder") from exc
    return {"path": data.get("path")}


@router.get("/threads/{thread_id}/dreamy/mount-folder/files")
async def list_mount_folder_files(thread_id: str) -> dict:
    """List files recursively inside the mounted folder for a thread."""
    config_path = _mount_config_path(thread_id)
    if not config_path.exists():
        return {"files": [], "folder_path": None}
    try:
        data = json.loads(config_path.read_text(encoding="utf-8"))
        folder_path_str = data.get("path")
    except Exception as exc:
        logger.error("Failed to read dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to read mounted folder config") from exc

    if not folder_path_str:
        return {"files": [], "folder_path": None}

    folder = Path(folder_path_str)
    if not folder.exists() or not folder.is_dir():
        return {"files": [], "folder_path": folder_path_str}

    try:
        files = []
        for p in sorted(folder.rglob("*")):
            rel = p.relative_to(folder).as_posix()
            if p.is_dir():
                files.append({
                    "name": f"{rel}/",
                    "size": 0,
                    "virtual_path": f"/mnt/user-data/mounted/{rel}",
                    "full_path": str(p),
                    "is_dir": True,
                })
                continue
            stat = p.stat()
            files.append({
                "name": rel,
                "size": stat.st_size,
                "virtual_path": f"/mnt/user-data/mounted/{rel}",
                "full_path": str(p),
                "is_dir": False,
            })
        return {"files": files, "folder_path": folder_path_str}
    except Exception as exc:
        logger.error("Failed to list mounted folder for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to list mounted folder files") from exc


@router.put("/threads/{thread_id}/dreamy/mount-folder")
async def put_mount_folder(thread_id: str, req: MountFolderRequest) -> dict:
    raw = (req.path or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="Path is required")
    if raw.startswith("/mnt/user-data/"):
        normalized_path = raw
    else:
        folder = Path(raw).expanduser().resolve()
        if not folder.exists() or not folder.is_dir():
            raise HTTPException(status_code=400, detail="Folder does not exist or is not a directory")
        normalized_path = str(folder)
    config_path = _mount_config_path(thread_id)
    config_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {"path": normalized_path}
    try:
        config_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as exc:
        logger.error("Failed to write dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to persist mounted folder") from exc
    return payload


@router.delete("/threads/{thread_id}/dreamy/mount-folder")
async def delete_mount_folder(thread_id: str) -> dict:
    config_path = _mount_config_path(thread_id)
    if not config_path.exists():
        return {"path": None}
    try:
        config_path.unlink()
    except Exception as exc:
        logger.error("Failed to delete dreamy mount config for thread %s: %s", thread_id, exc)
        raise HTTPException(status_code=500, detail="Failed to clear mounted folder") from exc
    return {"path": None}


# ─── Dreamy Executor control ──────────────────────────────────────────────────

def _signal_path(thread_id: str) -> Path:
    return get_paths().sandbox_outputs_dir(thread_id) / "pause_signal.json"


def _progress_path(thread_id: str) -> Path:
    return get_paths().sandbox_outputs_dir(thread_id) / "progress.json"


def _write_executor_signal(thread_id: str, signal: str) -> None:
    path = _signal_path(thread_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps({"signal": signal}), encoding="utf-8")


def _read_executor_progress(thread_id: str) -> dict:
    path = _progress_path(thread_id)
    if not path.exists():
        return {"state": "not_started"}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {"state": "unknown"}


@router.post("/threads/{thread_id}/dreamy/executor/pause")
async def pause_executor(thread_id: str) -> dict:
    """Signal the Dreamy Executor to stop at the next row boundary."""
    _write_executor_signal(thread_id, "pause")
    return {"signal": "pause", "thread_id": thread_id}


@router.post("/threads/{thread_id}/dreamy/executor/stop")
async def stop_executor(thread_id: str) -> dict:
    """Signal the Dreamy Executor to stop immediately after the current tool call."""
    _write_executor_signal(thread_id, "stop")
    return {"signal": "stop", "thread_id": thread_id}


@router.get("/threads/{thread_id}/dreamy/executor/status")
async def executor_status(thread_id: str) -> dict:
    """Return the current progress.json for the Dreamy Executor."""
    return _read_executor_progress(thread_id)


# ─── Native folder picker ─────────────────────────────────────────────────────

@router.get("/dreamy/pick-folder")
async def pick_folder() -> dict:
    """Open a native OS folder picker dialog and return the selected absolute path.

    macOS: uses osascript (always available).
    Linux: tries zenity, then kdialog, then yad.
    Returns {"path": "<absolute path>"} or {"path": null, "cancelled": true}.
    """
    import platform
    import subprocess as sp

    system = platform.system()
    try:
        if system == "Darwin":
            result = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: sp.run(
                    ["osascript", "-e", "POSIX path of (choose folder)"],
                    capture_output=True, text=True, timeout=120,
                ),
            )
            if result.returncode != 0:
                return {"path": None, "cancelled": True}
            path = result.stdout.strip().rstrip("/")
            return {"path": path, "cancelled": False}

        if system == "Linux":
            pickers = [
                ["zenity", "--file-selection", "--directory", "--title=Select Folder"],
                ["kdialog", "--getexistingdirectory", "/"],
                ["yad", "--file", "--directory", "--title=Select Folder"],
            ]
            for cmd in pickers:
                try:
                    result = await asyncio.get_event_loop().run_in_executor(
                        None,
                        lambda c=cmd: sp.run(c, capture_output=True, text=True, timeout=120),
                    )
                    if result.returncode == 0:
                        return {"path": result.stdout.strip().rstrip("/"), "cancelled": False}
                    # Return code 1 typically means user cancelled
                    return {"path": None, "cancelled": True}
                except FileNotFoundError:
                    continue
            raise HTTPException(status_code=501, detail="No folder picker available — install zenity or kdialog")

        raise HTTPException(status_code=501, detail=f"Native folder picker not supported on {system}")

    except sp.TimeoutExpired:
        return {"path": None, "cancelled": True}


# ─── macOS file actions ───────────────────────────────────────────────────────

class FileActionRequest(BaseModel):
    path: str


def _require_macos() -> None:
    if sys.platform != "darwin":
        raise HTTPException(status_code=400, detail="This action is only supported on macOS")


def _resolve_safe_path(raw: str) -> Path:
    """Resolve and lightly validate a path — must be absolute and must exist."""
    p = Path(raw).expanduser().resolve()
    if not p.exists():
        raise HTTPException(status_code=404, detail=f"Path not found: {raw}")
    return p


@router.post("/threads/{thread_id}/files/reveal")
async def reveal_in_finder(thread_id: str, req: FileActionRequest) -> dict:
    """Open the file's parent folder in Finder with the file selected (macOS only)."""
    _require_macos()
    p = _resolve_safe_path(req.path)
    asyncio.get_event_loop().run_in_executor(None, lambda: __import__("subprocess").Popen(["open", "-R", str(p)]))
    return {"success": True}


@router.post("/threads/{thread_id}/files/open")
async def open_in_default_app(thread_id: str, req: FileActionRequest) -> dict:
    """Open the file with its default application (macOS only)."""
    _require_macos()
    p = _resolve_safe_path(req.path)
    asyncio.get_event_loop().run_in_executor(None, lambda: __import__("subprocess").Popen(["open", str(p)]))
    return {"success": True}


@router.get("/threads/{thread_id}/files/thumbnail")
async def get_file_thumbnail(
    thread_id: str,
    path: str = Query(..., description="Absolute file path to generate thumbnail for"),
) -> Response:
    """Generate a Quick Look thumbnail for a file and return it as PNG (macOS only)."""
    _require_macos()
    p = _resolve_safe_path(path)

    with tempfile.TemporaryDirectory() as tmp_dir:
        proc = await asyncio.create_subprocess_exec(
            "qlmanage", "-t", "-s", "256", "-o", tmp_dir, str(p),
            stdout=asyncio.subprocess.DEVNULL,
            stderr=asyncio.subprocess.DEVNULL,
        )
        try:
            await asyncio.wait_for(proc.wait(), timeout=10.0)
        except TimeoutError:
            proc.kill()
            raise HTTPException(status_code=504, detail="Thumbnail generation timed out")

        thumbnails = list(Path(tmp_dir).glob("*.png"))
        if not thumbnails:
            raise HTTPException(status_code=404, detail="No thumbnail could be generated for this file type")

        return Response(content=thumbnails[0].read_bytes(), media_type="image/png")
